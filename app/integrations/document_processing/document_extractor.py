"""
Multi-Format Document Extraction Service

Service for extracting text content from multiple document formats:
- PDF (.pdf)
- Microsoft Word (.docx)
- Plain text (.txt)
- Markdown (.md)

Follows SRP: Single responsibility for document text extraction.
"""

import io
import logging
from pathlib import Path
from typing import Any

from .pdf_extractor import PDFExtractor

logger = logging.getLogger(__name__)

# Check for optional dependencies
try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None
    logger.warning("python-docx not installed. Install with: pip install python-docx")


class DocumentExtractor:
    """
    Service for extracting text from multiple document formats.

    Supported formats:
    - PDF (.pdf) - Uses PDFExtractor
    - Word (.docx) - Uses python-docx
    - Text (.txt) - Direct UTF-8 reading
    - Markdown (.md) - Direct UTF-8 reading

    Responsibilities:
    - Detect file format from filename
    - Extract text content from supported formats
    - Extract metadata when available
    - Handle errors gracefully

    Does NOT contain business logic (that's in Use Cases).
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    def __init__(self):
        """Initialize document extractor with format-specific extractors."""
        self._pdf_extractor = PDFExtractor()

    def extract(self, file_bytes: bytes, filename: str) -> dict[str, Any]:
        """
        Extract text and metadata from document bytes.

        Args:
            file_bytes: Document file content as bytes
            filename: Original filename (used to detect format)

        Returns:
            Dictionary with:
                - text: Full extracted text
                - metadata: Document metadata (format-specific)
                - page_count: Number of pages (for PDF) or None
                - source_filename: Original filename

        Raises:
            ValueError: If format is not supported or extraction fails
        """
        ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        # Route to appropriate extractor
        if ext == ".pdf":
            return self._extract_pdf(file_bytes, filename)
        elif ext == ".docx":
            return self._extract_docx(file_bytes, filename)
        elif ext in {".txt", ".md"}:
            return self._extract_text(file_bytes, filename)
        else:
            raise ValueError(f"No extractor available for format: {ext}")

    def _extract_pdf(self, file_bytes: bytes, filename: str) -> dict[str, Any]:
        """Extract text from PDF file."""
        try:
            result = self._pdf_extractor.extract_text_from_bytes(file_bytes)

            return {
                "text": result["text"],
                "metadata": result.get("metadata", {}),
                "page_count": result.get("page_count"),
                "pages": result.get("pages"),
                "source_filename": filename,
                "format": "pdf",
            }
        except Exception as e:
            logger.error(f"Error extracting PDF {filename}: {e}")
            raise ValueError(f"Could not extract text from PDF: {str(e)}") from e

    def _extract_docx(self, file_bytes: bytes, filename: str) -> dict[str, Any]:
        """Extract text from Word document."""
        if not DOCX_AVAILABLE or DocxDocument is None:
            raise ValueError(
                "python-docx is required for DOCX extraction. "
                "Install with: pip install python-docx"
            )

        try:
            doc_file = io.BytesIO(file_bytes)
            doc = DocxDocument(doc_file)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            full_text = "\n\n".join(paragraphs)

            # Extract metadata from core properties
            metadata = {}
            if doc.core_properties:
                props = doc.core_properties
                metadata = {
                    "title": props.title or "",
                    "author": props.author or "",
                    "subject": props.subject or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                }

            logger.info(
                f"Extracted {len(full_text)} characters from DOCX with "
                f"{len(paragraphs)} paragraphs"
            )

            return {
                "text": full_text,
                "metadata": metadata,
                "page_count": None,  # DOCX doesn't have fixed pages
                "paragraph_count": len(paragraphs),
                "source_filename": filename,
                "format": "docx",
            }

        except Exception as e:
            logger.error(f"Error extracting DOCX {filename}: {e}")
            raise ValueError(f"Could not extract text from DOCX: {str(e)}") from e

    def _extract_text(self, file_bytes: bytes, filename: str) -> dict[str, Any]:
        """Extract text from plain text or markdown file."""
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                text = file_bytes.decode("latin-1")
                logger.warning(f"File {filename} decoded with latin-1 fallback")

            # Clean up the text
            text = text.strip()

            # Count lines for metadata
            lines = text.split("\n")
            line_count = len(lines)

            ext = Path(filename).suffix.lower()
            format_type = "markdown" if ext == ".md" else "text"

            logger.info(f"Extracted {len(text)} characters from {format_type} file")

            return {
                "text": text,
                "metadata": {
                    "line_count": line_count,
                },
                "page_count": None,
                "source_filename": filename,
                "format": format_type,
            }

        except Exception as e:
            logger.error(f"Error extracting text file {filename}: {e}")
            raise ValueError(f"Could not extract text from file: {str(e)}") from e

    def is_supported(self, filename: str) -> bool:
        """
        Check if a filename has a supported extension.

        Args:
            filename: Filename to check

        Returns:
            True if format is supported, False otherwise
        """
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS

    def get_supported_extensions(self) -> set[str]:
        """Get set of supported file extensions."""
        return self.SUPPORTED_EXTENSIONS.copy()

    def validate_file(self, file_bytes: bytes, filename: str) -> bool:
        """
        Validate if a file can be extracted.

        Args:
            file_bytes: File content as bytes
            filename: Original filename

        Returns:
            True if file can be extracted, False otherwise
        """
        ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            return False

        if ext == ".pdf":
            return self._pdf_extractor.validate_pdf(file_bytes)
        elif ext == ".docx":
            if not DOCX_AVAILABLE or DocxDocument is None:
                return False
            try:
                doc_file = io.BytesIO(file_bytes)
                DocxDocument(doc_file)
                return True
            except Exception:
                return False
        elif ext in {".txt", ".md"}:
            # Text files are always valid if they can be decoded
            try:
                file_bytes.decode("utf-8")
                return True
            except UnicodeDecodeError:
                try:
                    file_bytes.decode("latin-1")
                    return True
                except Exception:
                    return False
        return False
